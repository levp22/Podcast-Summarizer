import os
from dotenv import load_dotenv
import openai
load_dotenv()
openai.api_key = os.getenv("OPEN_AI_KEY")
from docx import Document
# Path to your 'transcripts' directory
from datetime import datetime

# Iterate over each file in the directory
def create_summaries(transcript, gpt):
    with open('example.txt', 'r') as file:
        combined_message = file.read()
    # Now, file_content contains the entire content of 'yourfile.txt'
    if (gpt==1):
        part = openai.chat.completions.create(model="gpt-4-turbo-preview", 
            messages=[
            {"role": "system", "content": "You are an investment analyst who has been tasked with summarizing podcast transcripts for your boss so he can make strategic investment decisions."},
            {"role": "user", "content": "Please summarize the following podcast transcript for your boss in as much detail as possible. you must focus on the key points discussed, including the main arguments, insights, and takeaways. Additionally, identify any interesting quotes or counter-arguments presented. Do not miss a single point discussed in the transcript"},
            {"role": "user", "content": transcript},
            {"role": "user", "content": "there should be two sections in the output, the high level takeaways of each subject which includes key quotes, and then a detailed breakdown section as noted by the next prompt which dictates the schema of how the answer must be, try to stay away from generally saying what was talked about and instead focus on the specific details mentioned with some quotes and further information"},
    #  {"role": "user", "content": "it should include a section of high level key takeaways, some of which are crucial quotes but also include some specific takeaways from the whole thing - approximately 10 of them, then a section of a deep breakdown of the entire transcript that includes a header for each subject that is talked about in the podcast and include 5 detailed bullets for each subject, try to stay away from generally saying what was talked about and instead focus on the specific details mentioned with some quotes and further information"},
            {"role": "user", "content": combined_message},
            {"role": "user", "content": "you must ensure all topics are covered comprehensively according to the schema."}
        ]
        )
        #print(first_part)
        #print("part done")
        final_string = part.choices[0].message.content
    else:
        string_length = len(transcript)
        multiple = string_length // 30000 + 1
        div = string_length // multiple + multiple
        print(div)
        print(multiple)
        first_part = transcript[:div]
        second_part = transcript[div:]
        full_string = ""
        for x in range(1, multiple+1): 
            print(x)
            #Split the string into two halves
            part = openai.chat.completions.create(model="gpt-3.5-turbo", 
                messages=[
                {"role": "system", "content": "You are an investment analyst who has been tasked with summarizing podcast transcripts for your boss so he can make strategic investment insights."},
                {"role": "user", "content": "Please summarize the following podcast transcript for your boss in as much detail as possible. you must focus on the key points discussed, including the main arguments, insights, and takeaways. Additionally, identify any interesting quotes or counter-arguments presented. Please present in very detailed bullet point form, do not miss a single point discussed in the transcript"},
                {"role": "user", "content": first_part},
                {"role": "user", "content": "there should be two sections in the output, the high level takeaways of each subject which includes key quotes, and then a detailed breakdown section as noted by the next prompt, try to stay away from generally saying what was talked about and instead focus on the specific details mentioned with some quotes and further information"},
            #  {"role": "user", "content": "it should include a section of high level key takeaways, some of which are crucial quotes but also include some specific takeaways from the whole thing - approximately 10 of them, then a section of a deep breakdown of the entire transcript that includes a header for each subject that is talked about in the podcast and include 5 detailed bullets for each subject, try to stay away from generally saying what was talked about and instead focus on the specific details mentioned with some quotes and further information"},
                {"role": "user", "content": combined_message},
                {"role": "user", "content": "you must ensure all topics are covered comprehensively according to the schema."}
            ]
            )
            full_string = full_string + "\n" + part.choices[0].message.content
            #print(part.choices[0].message.content)
            if (x!=1):
                new_part = openai.chat.completions.create(model="gpt-3.5-turbo", 
                messages=[
                {"role": "system", "content": "You are an investment analyst who has been tasked with summarizing podcast transcripts for your boss so he can make strategic investment insights."},
                {"role": "user", "content": "Please combine the following two parts of a podcast summary into a single, cohesive document. You must ensure that the final summary is structured according to the specified schema with all relevant details maintained and properly integrated. The schema must be as follows:"},
                {"role": "user", "content": combined_message},
                {"role": "user", "content": "Here are the parts to combine, they are in the schema format as well:"},
            #  {"role": "user", "content": "it should include a section of high level key takeaways, some of which are crucial quotes but also include some specific takeaways from the whole thing - approximately 10 of them, then a section of a deep breakdown of the entire transcript that includes a header for each subject that is talked about in the podcast and include 5 detailed bullets for each subject, try to stay away from generally saying what was talked about and instead focus on the specific details mentioned with some quotes and further information"},
                {"role": "user", "content": "part 1: " + part.choices[0].message.content},
                {"role": "user", "content": "part 2: " + final_string},
                {"role": "user", "content": "you must ensure all topics are covered comprehensively according to the schema and that you maintain the two components of key highlights and detailed breakdown."}
                ]
                )
            #print(first_part)
            #print("part done")
                final_string = new_part.choices[0].message.content
            else:
                final_string = part.choices[0].message.content
                first_part = second_part[:div]
                second_part = second_part[div:]
    summary = openai.chat.completions.create(model="gpt-3.5-turbo", 
                messages=[
                    {"role": "system", "content": "You are an investment analyst who has been tasked with summarizing podcast transcripts for your boss so he can make strategic investment insights."},
                    {"role": "user", "content": "Write a 100 word summary of this summary of a podcast"},
                    {"role": "user", "content": final_string}
                ]
    )
    final_string = "Summary: " + summary.choices[0].message.content + "\n" + final_string 
    return final_string

#create_summaries({'What are Stocks and Bonds': {'author': 'Financeable Training'}, 'Buy Side vs Sell Side Firms': {'author': 'Financeable Training'}, 'Private Equity vs Venture Capital (and Growth Equity)': {'author': 'Financeable Training'}, 'What is an Investment Bank': {'author': 'Financeable Training'}, 'Sell Side M&A Process in Plain English': {'author': 'Financeable Training'}})
def create_word_documents():
    # Define default directories
    base_dir1 = 'key info'
    base_dir2 = 'summaries'
    output_dir = 'word files'

    print("Starting to create Word documents...")

    # Ensure the output directory exists, create if not
    os.makedirs(output_dir, exist_ok=True)
    print(f"Output directory set at: {output_dir}")

    # Count the number of documents created
    documents_created = 0

    # Walk through the directories in base_dir1
    for root, dirs, files in os.walk(base_dir1):
        print(f"Checking directory: {root}")
        for filename in files:
            # Construct the path for the corresponding file in base_dir2
            corresponding_path = os.path.join(root.replace(base_dir1, base_dir2), filename)

            # Check if the same file exists in base_dir2
            if os.path.exists(corresponding_path):
                print(f"Matching file found: {filename}")
                # Read content from both files
                with open(os.path.join(root, filename), 'r', encoding='utf-8') as file1:
                    content1 = file1.read()
                with open(corresponding_path, 'r', encoding='utf-8') as file2:
                    content2 = file2.read()
                
                # Create a new Word document
                doc = Document()
                doc.add_paragraph(content1)  # Insert content from first file
                doc.add_paragraph(content2)  # Insert content from second file

                # Define the output path for the Word document
                output_file_name = filename.replace('.txt', '.docx')
                output_path = os.path.join(output_dir, output_file_name)
                doc.save(output_path)  # Save the document
                print(f"Created Word document: {output_path}")
                documents_created += 1
    base_dir1 = 'key info'
    base_dir2 = 'unorganized summaries'
    output_dir = 'unorganized word files'

    print("Starting to create Word documents...")

    # Ensure the output directory exists, create if not
    os.makedirs(output_dir, exist_ok=True)
    print(f"Output directory set at: {output_dir}")

    # Count the number of documents created
    documents_created = 0

    # Walk through the directories in base_dir1
    for root, dirs, files in os.walk(base_dir1):
        print(f"Checking directory: {root}")
        for filename in files:
            # Construct the path for the corresponding file in base_dir2
            corresponding_path = os.path.join(root.replace(base_dir1, base_dir2), filename)

            # Check if the same file exists in base_dir2
            if os.path.exists(corresponding_path):
                print(f"Matching file found: {filename}")
                # Read content from both files
                with open(os.path.join(root, filename), 'r', encoding='utf-8') as file1:
                    content1 = file1.read()
                with open(corresponding_path, 'r', encoding='utf-8') as file2:
                    content2 = file2.read()
                
                # Create a new Word document
                doc = Document()
                doc.add_paragraph(content1)  # Insert content from first file
                doc.add_paragraph(content2)  # Insert content from second file

                # Define the output path for the Word document
                output_file_name = filename.replace('.txt', '.docx')
                output_file_name = "unorganized " + output_file_name
                output_path = os.path.join(output_dir, output_file_name)
                doc.save(output_path)  # Save the document
                print(f"Created Word document: {output_path}")
                documents_created += 1

    print(f"Process completed. Total documents created: {documents_created}")

# Execute the function

from pytube import YouTube
from pytube import Playlist
from youtube_transcript_api import YouTubeTranscriptApi
from datetime import date
import glob
from datetime import datetime
def get_todays_date():
    # Get today's date
    today = datetime.now()
    # Format the date as a string (e.g., "2024-03-14")
    date_string = today.strftime("%Y-%m-%d")
    return date_string
import subprocess
import sys
def create_audio_files():
    audio_files = os.fsencode("audio_files")
    for file in os.listdir(audio_files):
        filename = os.fsdecode(file)
    if filename.endswith(".mp4"):
        title = filename.replace(".mp4", "")

        # convert mp4 file to wav (16-bit)
        new_filename = title + ".wav"
        print("ran: "+ title)
        subprocess.run([
            'ffmpeg',
            '-i', os.path.join("audio_files", filename),
            '-ar', '16000',
            '-ac', '1',
            '-c:a', 'pcm_s16le',
            os.path.join("audio_files", new_filename)
        ])

        # call local whisper model
        result = subprocess.check_output([
            './main',
            '-f', '../Podcast-Summarizer/audio_files/' + new_filename
        ], cwd="../whisper.cpp")

        return result.decode(sys.stdout.encoding).strip()



def create_video_transcripts(yt):
    # Get video transcription, if none exist, download audio (last stream is guaranteed English)
    try:
        caption=YouTubeTranscriptApi.get_transcript(yt.video_id, languages=['en'])
        transcript = ""
        for obj in caption:
            transcript = transcript + obj['text']
        return transcript
    except:
        streams = yt.streams.filter(only_audio=True, file_extension="mp4")
        streams[-1].download("audio_files")
        return create_audio_files()


    
        
        
   ## description = yt.description

def create_key_info(yt):
    if yt.description is not None:
        return "Title:" + yt.title  + "\nAuthor: Lev Pollock\nDate: "+ get_todays_date() + "\nChannel: " + yt.author + "\nDescription: "+ yt.description + "\n"
    else:
        return "Title:" + yt.title  + "\nAuthor: Lev Pollock\nDate: "+ get_todays_date() + "\nChannel: " + yt.author+ "\n"

def create_transcripts_new(link,gpt):
    yt = YouTube(link)
    info = create_key_info(yt)
    transcript = create_video_transcripts(yt)
    summary = create_summaries(transcript, gpt)
    return info + "\n" + summary

def create_transcripts(type, dated):
    video_obj = {}
    if (type == 1):
        videos = open("inputs.txt", "r")
        directory_path = 'Transcripts'

        # Get a list of all .txt file paths in the directory
        txt_file_paths = glob.glob(os.path.join(directory_path, '*.txt'))

        # Loop through the list and remove each .txt file
        for file_path in txt_file_paths:
            try:
                os.remove(file_path)
                print(f"Deleted {file_path}")
            except Exception as e:
                print(f"Error deleting {file_path}: {e}")
        directory_path = 'audio_files'

        # Get a list of all .txt file paths in the directory
        txt_file_paths = glob.glob(os.path.join(directory_path, '*.mp4'))

        # Loop through the list and remove each .txt file
        for file_path in txt_file_paths:
            try:
                os.remove(file_path)
                print(f"Deleted {file_path}")
            except Exception as e:
                print(f"Error deleting {file_path}: {e}")
        txt_file_paths = glob.glob(os.path.join(directory_path, '*.wav'))

        # Loop through the list and remove each .txt file
        for file_path in txt_file_paths:
            try:
                os.remove(file_path)
                print(f"Deleted {file_path}")
            except Exception as e:
                print(f"Error deleting {file_path}: {e}")
        for video in videos:
            try:
                tit=create_video_transcripts(video)
                video_obj[tit] = { "author": yt.author}
                # Get YouTube video
                create_key_info(yt)
            except:
                print("Failed")
        create_audio_files()
        print(video_obj)
        create_summaries(video_obj)
    else :
        playlist = open("channels.txt", "r")
        directory_path = 'Transcripts'

        # Get a list of all .txt file paths in the directory
        txt_file_paths = glob.glob(os.path.join(directory_path, '*.txt'))

        # Loop through the list and remove each .txt file
        for file_path in txt_file_paths:
            try:
                os.remove(file_path)
                print(f"Deleted {file_path}")
            except Exception as e:
                print(f"Error deleting {file_path}: {e}")
        for channel in playlist:
            print(channel)
            c = Playlist(channel)
            for video in c.video_urls:
                try:
                    tit=create_video_transcripts(video)
                    yt = YouTube(video)
                    video_obj[tit] = { "author": yt.author}
                    # Get YouTube video
                    create_key_info(yt)
                except:
                    print("failed for :" + video)
        create_audio_files()
        print(video_obj)
        create_summaries(video_obj)
def playlist_dealer(link, gpt):
    strings = []
    c = Playlist(link)
    for video in c.video_urls:
        strings.append(create_transcripts_new(video, gpt))
    return strings


